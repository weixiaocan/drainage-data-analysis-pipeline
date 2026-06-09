"""测试 report_assembler 模块"""
import logging
import sys
from pathlib import Path

import pytest
import pandas as pd
from docx import Document

PROJECT_ROOT = Path(__file__).parent.parent


class TestReportAssembler:
    """报告组装模块测试"""

    @pytest.fixture
    def template_file(self) -> Path:
        return PROJECT_ROOT / "templates" / "监测数据分析报告模板-更新.docx"

    @pytest.fixture
    def combined_xlsx(self) -> Path:
        return PROJECT_ROOT / "outputs" / "综合分析结果.xlsx"

    @pytest.fixture
    def site_info_file(self) -> Path:
        return PROJECT_ROOT / "data" / "点位信息.xlsx"

    @pytest.fixture
    def output_file(self, tmp_path: Path) -> Path:
        return tmp_path / "测试报告.docx"

    def test_run_report_assembler_import(self):
        """测试模块可导入"""
        from src.pipeline.report_assembler.assembler import run_report_assembler
        assert callable(run_report_assembler)

    def test_run_report_assembler_runs(
        self,
        template_file: Path,
        combined_xlsx: Path,
        site_info_file: Path,
        output_file: Path,
    ):
        """测试运行不报错"""
        from src.pipeline.report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            combined_xlsx=combined_xlsx,
            site_info_file=site_info_file,
            output_file=output_file,
            has_rainfall_data=True,
        )

        assert result is not None

    def test_run_report_assembler_returns_dict(
        self,
        template_file: Path,
        combined_xlsx: Path,
        site_info_file: Path,
        output_file: Path,
    ):
        """测试返回结果格式"""
        from src.pipeline.report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            combined_xlsx=combined_xlsx,
            site_info_file=site_info_file,
            output_file=output_file,
            has_rainfall_data=True,
        )

        assert "output_file" in result
        assert "stats" in result
        assert isinstance(result["output_file"], Path)
        assert isinstance(result["stats"], dict)

    def test_output_file_created(
        self,
        template_file: Path,
        combined_xlsx: Path,
        site_info_file: Path,
        output_file: Path,
    ):
        """测试输出文件创建"""
        from src.pipeline.report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            combined_xlsx=combined_xlsx,
            site_info_file=site_info_file,
            output_file=output_file,
            has_rainfall_data=True,
        )

        assert result["output_file"].exists()

    def test_stats_structure(
        self,
        template_file: Path,
        combined_xlsx: Path,
        site_info_file: Path,
        output_file: Path,
    ):
        """测试统计信息结构"""
        from src.pipeline.report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            combined_xlsx=combined_xlsx,
            site_info_file=site_info_file,
            output_file=output_file,
            has_rainfall_data=True,
        )

        stats = result["stats"]
        assert "tables_filled" in stats
        assert "images_inserted" in stats
        assert "points_processed" in stats

    def test_template_scanner_detects_key_roles(self, template_file: Path):
        """当前 Word 模板应能被识别为语义表格，而不是依赖固定下标。"""
        from src.pipeline.report_assembler.template_scanner import scan_template

        mapping = scan_template(Document(template_file))

        for role in [
            "site_info",
            "collection_rate",
            "rainfall_daily",
            "rainfall_events",
            "dry_risk",
            "rainy_overflow_risk",
        ]:
            assert mapping.get(role) is not None
        assert len(mapping.curve_tables) >= 1

    def test_legacy_template_point_name_keeps_full_point_id(self):
        """旧模板编号只去掉轮次前缀，不能截断复合点位名。"""
        from src.pipeline.report_assembler.data_context import clean_point_id

        assert clean_point_id("1-1#") == "#1"
        assert clean_point_id("1-2-3#") == "#2-3"
        assert clean_point_id("1-2-19-1#") == "#2-19-1"

    def test_dry_risk_table_uses_explicit_column_mapping(self, template_file: Path):
        """风险表不能受 DataFrame 原始列顺序影响。"""
        from src.pipeline.report_assembler.data_context import ReportDataContext
        from src.pipeline.report_assembler.report_tables import TABLE_SPECS, render_report_table
        from src.pipeline.report_assembler.template_scanner import scan_template

        doc = Document(template_file)
        mapping = scan_template(doc)
        table = mapping.get("dry_risk")
        df = pd.DataFrame(
            [
                {
                    "overflow_risk": "低溢流风险",
                    "running_risk": "中风险",
                    "silting_risk": "高淤积风险",
                    "overflow_value": 0.54,
                    "max_fullness": 1.95,
                    "max_level_m": 2.92,
                    "dry_velocity_mps": 0.04,
                    "well_depth_m": 5.4,
                    "diameter_m": 1.5,
                    "point_id": "#1",
                    "index": 1,
                }
            ]
        )
        context = ReportDataContext(
            analysis={"dry_risk": df},
            site_info=pd.DataFrame(),
            point_ids=["#1"],
        )

        render_report_table(table, TABLE_SPECS["dry_risk"], context)

        row = table.rows[2].cells
        assert row[1].text == "#1"
        assert row[4].text == "0.04"
        assert row[8].text == "高淤积风险"
        assert row[10].text == "低溢流风险"

    def test_output_point_names_do_not_use_template_prefix(
        self,
        template_file: Path,
        combined_xlsx: Path,
        site_info_file: Path,
        output_file: Path,
    ):
        """报告中的点位名应展示真实 point_id，不能出现旧模板 1- 前缀。"""
        from src.pipeline.report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            combined_xlsx=combined_xlsx,
            site_info_file=site_info_file,
            output_file=output_file,
            has_rainfall_data=True,
        )
        doc = Document(result["output_file"])
        from src.pipeline.report_assembler.template_scanner import scan_template

        mapping = scan_template(doc)
        checks = [
            ("site_info", 1, 0),
            ("collection_rate", 1, 0),
            ("dry_risk", 2, 1),
            ("rainy_overflow_risk", 1, 0),
        ]
        for role, first_data_row, cell_idx in checks:
            table = mapping.get(role)
            if table is None or len(table.rows) <= first_data_row:
                continue
            point_text = table.rows[first_data_row].cells[cell_idx].text
            assert not point_text.startswith("1-")

    def test_output_has_no_template_residue(
        self,
        template_file: Path,
        combined_xlsx: Path,
        site_info_file: Path,
        output_file: Path,
    ):
        """关键模板旧内容必须被替换。"""
        from src.pipeline.report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            combined_xlsx=combined_xlsx,
            site_info_file=site_info_file,
            output_file=output_file,
            has_rainfall_data=True,
        )
        text = "\n".join(p.text for p in Document(result["output_file"]).paragraphs)

        assert "2024/9/18" not in text
        assert "2024/11/26" not in text
        assert "13台流量监测设备" not in text
        assert "1-9#" not in text


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
