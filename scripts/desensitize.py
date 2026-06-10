"""
数据脱敏工具

功能：
1. 时间偏移：年份+1，月份+1
2. 点位编号统一重命名
3. 地名替换
"""

import os
import re
import pandas as pd
from datetime import datetime, timedelta
from docx import Document


# ============================================================
# 点位编号映射
# ============================================================

# 原始点位编号列表（按出现顺序）
ROUND1_POINTS = ['#1', '#5', '#7', '#9', '#15', '#16', '13']  # 第一轮
ROUND2_POINTS = ['#2-2', '2-3', '#2-4', '2-7', '#2-8', '#2-14', '2-12', '2-15', '#2-16', '#2-19', '#2-19-1', '#2-5']  # 第二轮

def build_point_mapping():
    """构建点位编号映射表"""
    mapping = {}

    # 第一轮点位：监测点1-01, 监测点1-02, ...
    for i, point in enumerate(ROUND1_POINTS, 1):
        new_id = f"监测点1-{i:02d}"
        # 处理带 # 和不带 # 的情况
        mapping[point] = new_id
        if point.startswith('#'):
            mapping[point[1:]] = new_id  # 也映射不带 # 的版本
        else:
            mapping[f'#{point}'] = new_id  # 也映射带 # 的版本

    # 第二轮点位：监测点2-01, 监测点2-02, ...
    for i, point in enumerate(ROUND2_POINTS, 1):
        new_id = f"监测点2-{i:02d}"
        mapping[point] = new_id
        if point.startswith('#'):
            mapping[point[1:]] = new_id
        else:
            mapping[f'#{point}'] = new_id

    return mapping

POINT_MAPPING = build_point_mapping()


def replace_point_id(text: str) -> str:
    """替换文本中的点位编号"""
    if not isinstance(text, str):
        return text

    result = text
    for old_id, new_id in POINT_MAPPING.items():
        # 替换各种可能的格式
        patterns = [
            old_id,
            old_id.replace('#', ''),
            f'#{old_id}' if not old_id.startswith('#') else old_id,
        ]
        for pattern in patterns:
            if pattern in result:
                result = result.replace(pattern, new_id)

    return result


def get_new_point_id(old_id: str) -> str:
    """获取点位的新编号"""
    # 标准化输入
    old_id = str(old_id).strip()
    if old_id in POINT_MAPPING:
        return POINT_MAPPING[old_id]
    # 尝试不带 # 的版本
    if old_id.startswith('#') and old_id[1:] in POINT_MAPPING:
        return POINT_MAPPING[old_id[1:]]
    # 尝试带 # 的版本
    if f'#{old_id}' in POINT_MAPPING:
        return POINT_MAPPING[f'#{old_id}']
    return old_id  # 未找到则返回原值


# ============================================================
# 时间偏移
# ============================================================

def shift_datetime(dt, year_offset=1, month_offset=1):
    """
    时间偏移：年份+1，月份+1

    Args:
        dt: datetime 对象或字符串
        year_offset: 年份偏移量
        month_offset: 月份偏移量

    Returns:
        偏移后的 datetime 对象
    """
    if isinstance(dt, str):
        # 尝试解析多种日期格式
        for fmt in ['%Y/%m/%d %H:%M', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d', '%Y/%m/%d']:
            try:
                dt = datetime.strptime(dt, fmt)
                break
            except ValueError:
                continue
        else:
            return dt  # 无法解析，返回原值

    if not isinstance(dt, datetime):
        return dt

    # 计算新年份和月份
    new_year = dt.year + year_offset
    new_month = dt.month + month_offset

    # 处理月份溢出
    while new_month > 12:
        new_month -= 12
        new_year += 1

    # 处理日期溢出（如 1月31日 + 1个月 → 2月28/29日）
    import calendar
    max_day = calendar.monthrange(new_year, new_month)[1]
    new_day = min(dt.day, max_day)

    return dt.replace(year=new_year, month=new_month, day=new_day)


def shift_date_string(date_str: str) -> str:
    """
    偏移日期字符串，保持原格式

    支持格式：
    - 2026/2/1 0:00 → 2027/3/1 0:00
    - 2026-03-07 00:00:00 → 2027-04-07 00:00:00
    - 2026年2月 → 2027年3月
    """
    if not isinstance(date_str, str):
        return date_str

    # 格式1: 2026/2/1 0:00
    match = re.match(r'(\d{4})/(\d{1,2})/(\d{1,2})\s*(\d{1,2}:\d{2})?', date_str)
    if match:
        year, month, day, time_part = match.groups()
        new_dt = shift_datetime(datetime(int(year), int(month), int(day)))
        time_str = f' {time_part}' if time_part else ''
        return f'{new_dt.year}/{new_dt.month}/{new_dt.day}{time_str}'

    # 格式2: 2026-03-07 00:00:00
    match = re.match(r'(\d{4})-(\d{2})-(\d{2})\s*(\d{2}:\d{2}:\d{2})?', date_str)
    if match:
        year, month, day, time_part = match.groups()
        new_dt = shift_datetime(datetime(int(year), int(month), int(day)))
        time_str = f' {time_part}' if time_part else ''
        return f'{new_dt.year}-{new_dt.month:02d}-{new_dt.day:02d}{time_str}'

    # 格式3: 2026年2月
    match = re.match(r'(\d{4})年(\d{1,2})月', date_str)
    if match:
        year, month = int(match.group(1)), int(match.group(2))
        new_dt = shift_datetime(datetime(year, month, 1))
        return f'{new_dt.year}年{new_dt.month}月'

    return date_str  # 无法匹配，返回原值


# ============================================================
# 地名替换
# ============================================================

PLACE_MAPPING = {
    '前川片区': '监测片区',
    '前川': '监测',
    '前川大道': '主干道A',
    '西寺大道': '道路A',
    '石阳街': '街道A',
    '向阳大街': '道路B',
    '鲁台泵站': '泵站A',
    '鲁台': '泵站',
    '黄陂': '区域',
}

def replace_place_names(text: str) -> str:
    """替换地名"""
    if not isinstance(text, str):
        return text

    result = text
    for old_name, new_name in PLACE_MAPPING.items():
        result = result.replace(old_name, new_name)

    return result


# ============================================================
# 文件处理函数
# ============================================================

def process_rainfall_data(input_path: str, output_path: str):
    """处理降雨数据 CSV 文件"""
    print(f"处理降雨数据: {input_path}")

    df = pd.read_csv(input_path, encoding='utf-8')

    # 偏移日期列
    if 'date' in df.columns:
        df['date'] = df['date'].apply(shift_date_string)

    df.to_csv(output_path, index=False, encoding='utf-8')
    print(f"  已保存: {output_path}")


def process_point_info(input_path: str, output_path: str):
    """处理点位信息 Excel 文件"""
    print(f"处理点位信息: {input_path}")

    df = pd.read_excel(input_path)

    # 获取第一列名（安装点位列）
    first_col = df.columns[0]

    # 替换点位编号
    df[first_col] = df[first_col].apply(lambda x: get_new_point_id(x))

    # 查找并偏移时间列
    for col in df.columns:
        if '时间' in col or '日期' in col:
            df[col] = df[col].apply(lambda x: shift_date_string(str(x)) if pd.notna(x) else x)

    df.to_excel(output_path, index=False)
    print(f"  已保存: {output_path}")


def process_flow_csv(input_path: str, output_path: str, old_point_id: str):
    """处理流量数据 CSV 文件"""
    print(f"处理流量数据: {input_path}")

    df = pd.read_csv(input_path, encoding='utf-8')

    # 偏移时间列
    if '数据时间' in df.columns:
        df['数据时间'] = df['数据时间'].apply(shift_date_string)

    df.to_csv(output_path, index=False, encoding='utf-8')
    print(f"  已保存: {output_path}")


def process_template(input_path: str, output_path: str):
    """处理 Word 模板文件"""
    print(f"处理模板文件: {input_path}")

    doc = Document(input_path)

    # 替换段落中的文本
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_place_names(run.text)
            run.text = replace_point_id(run.text)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = replace_place_names(run.text)
                        run.text = replace_point_id(run.text)

    doc.save(output_path)
    print(f"  已保存: {output_path}")


# ============================================================
# 主函数
# ============================================================

def main():
    """执行完整脱敏流程"""
    import shutil

    # 定义路径
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    data_dir = os.path.join(base_dir, 'data')
    flow_dir = os.path.join(data_dir, 'flow')
    template_dir = os.path.join(base_dir, 'templates')

    print("=" * 60)
    print("开始数据脱敏")
    print("=" * 60)

    # 1. 处理降雨数据
    rainfall_path = os.path.join(data_dir, '降雨数据.csv')
    if os.path.exists(rainfall_path):
        process_rainfall_data(rainfall_path, rainfall_path)

    # 2. 处理点位信息
    point_info_path = os.path.join(data_dir, '点位信息.xlsx')
    if os.path.exists(point_info_path):
        process_point_info(point_info_path, point_info_path)

    # 3. 处理流量数据文件
    if os.path.exists(flow_dir):
        # 建立文件名映射
        for filename in os.listdir(flow_dir):
            if not filename.endswith('.csv'):
                continue

            old_path = os.path.join(flow_dir, filename)

            # 从文件名提取点位编号 (格式: 35891_#1.csv)
            parts = filename.replace('.csv', '').split('_')
            if len(parts) >= 2:
                old_point_id = parts[1]
                new_point_id = get_new_point_id(old_point_id)

                # 新文件名保持设备ID前缀，替换点位编号
                new_filename = f"{parts[0]}_{new_point_id}.csv"
                new_path = os.path.join(flow_dir, new_filename)

                # 处理文件内容
                process_flow_csv(old_path, new_path, old_point_id)

                # 如果文件名改变了，删除旧文件
                if old_path != new_path and os.path.exists(old_path):
                    os.remove(old_path)
                    print(f"  删除旧文件: {old_path}")

    # 4. 处理模板文件
    template_path = os.path.join(template_dir, '监测数据分析报告模板-更新.docx')
    if os.path.exists(template_path):
        process_template(template_path, template_path)

    print("=" * 60)
    print("数据脱敏完成")
    print("=" * 60)


if __name__ == '__main__':
    main()
