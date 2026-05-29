"""报告组装核心逻辑

将各模块的分析结果组装到 Word 报告模板中。

主要步骤：
1. 读取综合分析结果.xlsx的各个sheet
2. 读取点位信息.xlsx
3. 动态调整表格行数
4. 执行规则文字替换
5. 生成排污规律章节
6. LLM生成风险分析段落
7. 插入特征曲线图
8. 保存最终报告
"""

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook

from .table_manager import adjust_table_rows, adjust_curve_image_tables, clear_table_data
from .text_replacer import TextReplacer, build_context_from_data
from .pattern_section import get_pattern_point_names, insert_curve_images_to_tables
from .llm_writer import LLMReportWriter


# ===== 点位编号转换函数 =====

def to_template_point_name(data_name: str) -> str:
    """数据编号 → 模板编号

    #1 → 1-1#
    """
    num = data_name.replace("#", "")
    return f"1-{num}#"


def to_data_point_name(template_name: str) -> str:
    """模板编号 → 数据编号

    1-1# → #1
    """
    parts = template_name.replace("#", "").split("-")
    return f"#{parts[-1]}"


@dataclass
class ReportConfig:
    """报告组装配置参数"""
    monitoring_start: str = ""
    monitoring_end: str = ""
    monitoring_round: str = "第一轮"
    rainfall_threshold_mm: float = 2.0


def _load_analysis_results(xlsx_path: Path) -> Dict[str, pd.DataFrame]:
    """加载综合分析结果.xlsx的所有sheet"""
    results = {}
    xlsx_path = Path(xlsx_path)
    if not xlsx_path.exists():
        print(f"警告: 分析结果文件不存在: {xlsx_path}")
        return results

    wb = load_workbook(xlsx_path)
    for sheet_name in wb.sheetnames:
        if sheet_name.startswith("特征曲线_"):
            continue

        ws = wb[sheet_name]
        data = []
        for row in ws.iter_rows(values_only=True):
            data.append(row)
        if data:
            df = pd.DataFrame(data[1:], columns=data[0])
            results[sheet_name] = df
    return results


def _load_site_info(xlsx_path: Path) -> pd.DataFrame:
    """加载点位信息.xlsx"""
    if not xlsx_path.exists():
        print(f"警告: 点位信息文件不存在: {xlsx_path}")
        return pd.DataFrame()

    wb = load_workbook(xlsx_path)
    ws = wb.active
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append(row)
    if data:
        return pd.DataFrame(data[1:], columns=data[0])
    return pd.DataFrame()


def _fill_table_with_df(table, df: pd.DataFrame, start_row: int = 1) -> None:
    """用DataFrame填充表格

    Args:
        table: docx表格对象
        df: 数据
        start_row: 开始填充的行（跳过表头）
    """
    for i, (_, row) in enumerate(df.iterrows()):
        if start_row + i >= len(table.rows):
            break
        table_row = table.rows[start_row + i]
        for j, value in enumerate(row):
            if j >= len(table_row.cells):
                break
            cell = table_row.cells[j]
            if pd.isna(value):
                cell.text = ""
            elif isinstance(value, (int, float)):
                cell.text = str(round(value, 2) if isinstance(value, float) else value)
            else:
                cell.text = str(value)


def _fill_site_info_table(table, site_info: pd.DataFrame, data_point_names: list) -> None:
    """填充监测点位安装信息表 (Table 0)"""
    site_map = {}
    for _, row in site_info.iterrows():
        point_code = str(row.get("监测点编号", ""))
        if point_code:
            site_map[point_code] = row

    for i, data_name in enumerate(data_point_names):
        if i + 1 >= len(table.rows):
            break

        template_name = to_template_point_name(data_name)
        row_data = None

        for key in [data_name, data_name.replace("#", ""), template_name]:
            if key in site_map:
                row_data = site_map[key]
                break

        if row_data is None:
            continue

        table_row = table.rows[i + 1]
        cells = table_row.cells

        if len(cells) > 0:
            cells[0].text = template_name
        if len(cells) > 1:
            cells[1].text = str(row_data.get("设备类型", ""))
        if len(cells) > 2:
            shape = row_data.get("绑定管形状", "")
            cells[2].text = "圆管" if shape == "圆管" else str(shape) if shape else ""
        if len(cells) > 3:
            cells[3].text = str(row_data.get("管径(m)", ""))
        if len(cells) > 4:
            cells[4].text = str(row_data.get("井深(m)", ""))
        if len(cells) > 5:
            install_time = row_data.get("设备安装时间", "")
            if install_time:
                cells[5].text = str(install_time)[:10]


def _fill_collection_rate_table(table, analysis_data: dict, data_point_names: list, filter_result_path: Path | None = None) -> None:
    """填充数据收集率统计表 (Table 1)"""
    collection_data = {}

    # 优先从综合分析结果读取
    if "数据收集率统计" in analysis_data:
        df = analysis_data["数据收集率统计"]
        for _, row in df.iterrows():
            point_name = str(row.get("点位编号", ""))
            # 提取点位编号（如 #1）
            if "#" in point_name:
                import re
                match = re.search(r'#\d+', point_name)
                if match:
                    simple_name = match.group()
                    collection_data[simple_name] = {
                        "count": row.get("监测数据条数", 0),
                        "days": row.get("监测天数", 0),
                        "theoretical": row.get("理论数据条数", 0),
                        "rate": row.get("数据收集率(%)", 0) / 100 if row.get("数据收集率(%)") else 0,
                    }

    for i, data_name in enumerate(data_point_names):
        if i + 1 >= len(table.rows):
            break

        table_row = table.rows[i + 1]
        cells = table_row.cells

        template_name = to_template_point_name(data_name)
        data = collection_data.get(data_name, {})

        if len(cells) > 0:
            cells[0].text = template_name
        if len(cells) > 1:
            cells[1].text = str(data.get("count", "-"))
        if len(cells) > 2:
            cells[2].text = str(data.get("days", "-"))
        if len(cells) > 3:
            cells[3].text = str(data.get("theoretical", "-"))
        if len(cells) > 4:
            rate = data.get("rate", 0)
            if rate > 0:
                cells[4].text = f"{rate * 100:.1f}%"
            else:
                cells[4].text = "-"


def _fill_rainfall_daily_table(table, df: pd.DataFrame) -> None:
    """填充日降雨量统计表 (Table 2)"""
    if df.empty or "日降雨量(mm)" not in df.columns:
        return

    df_rain = df[df["日降雨量(mm)"] > 0].copy()

    for i, (_, row) in enumerate(df_rain.iterrows()):
        if i + 1 >= len(table.rows):
            break

        table_row = table.rows[i + 1]
        cells = table_row.cells

        date_val = row.get("日期", "")
        if date_val:
            try:
                if hasattr(date_val, "strftime"):
                    date_str = date_val.strftime("%Y-%m-%d")
                else:
                    date_str = str(date_val)[:10]
            except Exception:
                date_str = str(date_val)[:10]

            if len(cells) > 0:
                cells[0].text = date_str

        if len(cells) > 1:
            rainfall = row.get("日降雨量(mm)", 0)
            cells[1].text = str(round(rainfall, 1) if isinstance(rainfall, float) else rainfall)


def _fill_rainfall_event_table(table, df: pd.DataFrame) -> None:
    """填充场次降雨统计表 (Table 3)"""
    if df.empty:
        return

    col_mapping = {
        "场次编号": "编号",
        "开始时间": "开始时间",
        "结束时间": "结束时间",
        "总降雨量(mm)": "总降雨量/mm",
        "降雨历时(h)": "降雨历时/h",
        "平均强度(mm/h)": "平均强度/mm/h",
        "降雨等级": "降雨等级",
    }

    for i, (_, row) in enumerate(df.iterrows()):
        if i + 1 >= len(table.rows):
            break

        table_row = table.rows[i + 1]
        cells = table_row.cells

        for j, (src_col, dst_name) in enumerate(col_mapping.items()):
            if j >= len(cells):
                break

            value = row.get(src_col, "")
            if pd.isna(value):
                cells[j].text = ""
            elif isinstance(value, float):
                cells[j].text = str(round(value, 2))
            elif hasattr(value, "strftime"):
                cells[j].text = value.strftime("%Y-%m-%d %H:%M")
            else:
                cells[j].text = str(value)


def _insert_curve_images(doc: Document, img_dir: Path, data_point_names: list) -> int:
    """插入特征曲线图到对应表格 (Tables 4-16)"""
    inserted_count = 0

    for i, data_name in enumerate(data_point_names):
        table_idx = 4 + i
        if table_idx >= len(doc.tables):
            break

        img_path = img_dir / f"{data_name}_特征曲线.png"
        if not img_path.exists():
            continue

        table = doc.tables[table_idx]
        if len(table.rows) == 0 or len(table.rows[0].cells) < 2:
            continue

        try:
            for cell_idx in [0, 1]:
                cell = table.rows[0].cells[cell_idx]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(str(img_path), width=Inches(2.8))
                inserted_count += 1
        except Exception as e:
            print(f"插入图片失败 {data_name}: {e}")

    return inserted_count


def _fill_risk_table(table, df: pd.DataFrame, start_row: int = 2) -> None:
    """填充旱天风险表格 (Table 17)"""
    if df.empty:
        return

    for i, (_, row) in enumerate(df.iterrows()):
        if start_row + i >= len(table.rows):
            break

        table_row = table.rows[start_row + i]
        cells = table_row.cells

        for j, value in enumerate(row):
            if j >= len(cells):
                break

            if pd.isna(value):
                cells[j].text = ""
            elif isinstance(value, float):
                cells[j].text = str(round(value, 2))
            else:
                cells[j].text = str(value)


def _generate_pattern_curve_image(curve: pd.DataFrame, point_name: str, output_dir: Path) -> Path | None:
    """生成排污规律特征曲线图"""
    try:
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates

        plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "SimSun"]
        plt.rcParams["axes.unicode_minus"] = False

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6), sharex=True)

        ax1.plot(curve.index, curve["f"], 'b-', linewidth=1)
        ax1.set_ylabel("流量 (L/s)")
        ax1.set_title(f"{point_name} 流量特征曲线")
        ax1.grid(True, alpha=0.3)

        ax2.plot(curve.index, curve["l"], 'g-', linewidth=1)
        ax2.set_ylabel("液位 (m)")
        ax2.set_title(f"{point_name} 液位特征曲线")
        ax2.grid(True, alpha=0.3)

        ax2.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M"))
        ax2.set_xlabel("时间")

        plt.tight_layout()

        output_dir.mkdir(parents=True, exist_ok=True)
        img_path = output_dir / f"{point_name}_特征曲线.png"
        plt.savefig(img_path, dpi=150, bbox_inches="tight")
        plt.close(fig)

        return img_path

    except Exception as e:
        print(f"生成特征曲线图失败: {e}")
        return None


def _remove_rainy_sections(doc: Document) -> None:
    """删除报告中的雨天相关章节"""
    rainy_keywords = ["降雨分析", "降雨情况", "雨天溢流", "雨天运行风险"]

    paragraphs_to_delete = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for keyword in rainy_keywords:
            if keyword in text:
                paragraphs_to_delete.append(i)
                print(f"  标记删除段落: {text[:50]}...")
                break

    for i in reversed(paragraphs_to_delete):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    tables = doc.tables
    rainy_table_indices = [2, 3, 18]

    for idx in rainy_table_indices:
        if idx < len(tables):
            table = tables[idx]
            for row in table.rows[1:]:
                for cell in row.cells:
                    cell.text = ""
            print(f"  清空表格{idx}内容")


def _find_paragraph_index(doc: Document, keyword: str) -> int:
    """通过关键词定位段落索引"""
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            return i
    return -1


def _replace_paragraph_content(para, new_text: str) -> None:
    """替换段落内容，保留格式"""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def _insert_paragraphs_after(doc: Document, after_idx: int, texts: list) -> None:
    """在指定段落后插入多个段落"""
    if after_idx < 0 or after_idx >= len(doc.paragraphs):
        return

    para = doc.paragraphs[after_idx]
    for text in reversed(texts):
        new_para = para.insert_paragraph_before(text)
        para._element.addnext(new_para._element)


def run_report_assembler(
    template_file: Path,
    combined_xlsx: Path,
    site_info_file: Path,
    output_file: Path,
    dry_curve_data: Dict[str, pd.DataFrame] | None = None,
    filter_result_path: Path | None = None,
    config: Dict[str, Any] | None = None,
    has_rainfall_data: bool = True,
    llm_client = None,
) -> Dict[str, Any]:
    """执行报告组装

    Args:
        template_file: Word报告模板文件
        combined_xlsx: 综合分析结果.xlsx
        site_info_file: 点位信息.xlsx
        output_file: 输出报告文件路径
        dry_curve_data: 旱天特征曲线数据（从内存传入，可选）
        filter_result_path: 筛选结果.xlsx路径
        config: 可选配置参数
        has_rainfall_data: 是否有降雨数据
        llm_client: LLM客户端（用于生成风险分析段落）

    Returns:
        {
            "output_file": Path,
            "stats": dict,
        }
    """
    cfg = ReportConfig()
    if config:
        for key, value in config.items():
            if hasattr(cfg, key):
                setattr(cfg, key, value)

    output_file.parent.mkdir(parents=True, exist_ok=True)

    # 1. 加载数据
    print(f"读取报告模板: {template_file}")
    doc = Document(template_file)

    print(f"读取综合分析结果: {combined_xlsx}")
    analysis_data = _load_analysis_results(combined_xlsx)

    print(f"读取点位信息: {site_info_file}")
    site_info = _load_site_info(site_info_file)

    # 如果没有传入 dry_curve_data，尝试从 Excel 读取
    if dry_curve_data is None:
        dry_curve_data = {}
        try:
            wb = load_workbook(combined_xlsx, data_only=True)
            for sheet_name in wb.sheetnames:
                if sheet_name.startswith("特征曲线_"):
                    ws = wb[sheet_name]
                    point_name = sheet_name.replace("特征曲线_", "")

                    data = []
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        if row[0] is not None:
                            data.append(row)

                    if data:
                        df = pd.DataFrame(data, columns=["时间", "流量(L/s)", "液位(m)", "流速(m/s)"])
                        df = df.dropna(subset=["时间"])
                        df["时间"] = pd.date_range("00:00:00", "23:59:00", freq="T")[:len(df)]
                        df = df.set_index("时间")
                        df = df.rename(columns={"流量(L/s)": "f", "液位(m)": "l", "流速(m/s)": "velo"})
                        dry_curve_data[point_name] = df
            wb.close()
        except Exception as e:
            print(f"读取旱天特征曲线数据失败: {e}")

    stats = {
        "tables_filled": 0,
        "images_inserted": 0,
        "points_processed": 0,
        "text_replaced": 0,
        "llm_generated": 0,
    }

    tables = doc.tables
    print(f"\n报告包含 {len(tables)} 个表格")

    # 获取点位编号列表
    data_point_names = list(dry_curve_data.keys())
    if not data_point_names and "旱天分析" in analysis_data:
        df = analysis_data["旱天分析"]
        if "点位编号" in df.columns:
            data_point_names = df["点位编号"].dropna().tolist()

    print(f"点位列表: {data_point_names}")

    # ========== 2. 表格行数动态调整 ==========
    print("\n动态调整表格行数...")

    # 表格0: 监测点位安装信息
    if len(tables) > 0 and data_point_names:
        target_rows = len(data_point_names)
        adjust_table_rows(tables[0], target_rows)
        print(f"  表格0: 调整为 {target_rows + 1} 行")

    # 表格1: 数据收集率统计
    if len(tables) > 1 and data_point_names:
        target_rows = len(data_point_names)
        adjust_table_rows(tables[1], target_rows)
        print(f"  表格1: 调整为 {target_rows + 1} 行")

    # 表格2: 日降雨量统计
    if has_rainfall_data and len(tables) > 2 and "日降雨量统计" in analysis_data:
        rainy_days = len(analysis_data["日降雨量统计"][analysis_data["日降雨量统计"]["日降雨量(mm)"] > 0])
        adjust_table_rows(tables[2], rainy_days)
        print(f"  表格2: 调整为 {rainy_days + 1} 行")

    # 表格3: 场次降雨统计
    if has_rainfall_data and len(tables) > 3 and "场次降雨统计" in analysis_data:
        event_count = len(analysis_data["场次降雨统计"])
        adjust_table_rows(tables[3], event_count)
        print(f"  表格3: 调整为 {event_count + 1} 行")

    # 表格17: 旱天风险
    if len(tables) > 17 and "旱天风险" in analysis_data:
        target_rows = len(analysis_data["旱天风险"])
        adjust_table_rows(tables[17], target_rows, template_row_idx=2)
        print(f"  表格17: 调整为 {target_rows + 2} 行")

    # 表格18: 雨天溢流风险
    if has_rainfall_data and len(tables) > 18 and "雨天溢流风险" in analysis_data:
        target_rows = len(analysis_data["雨天溢流风险"])
        adjust_table_rows(tables[18], target_rows)
        print(f"  表格18: 调整为 {target_rows + 1} 行")

    # ========== 3. 规则文字替换 ==========
    print("\n执行规则文字替换...")
    context = build_context_from_data(
        collection_df=analysis_data.get("数据收集率统计"),
        rainfall_daily_df=analysis_data.get("日降雨量统计"),
        rainfall_event_df=analysis_data.get("场次降雨统计"),
        site_info_df=site_info,
    )
    replacer = TextReplacer(context)
    replaced_count = replacer.replace_in_document(doc)
    stats["text_replaced"] = replaced_count
    print(f"  替换了 {replaced_count} 个段落")

    # ========== 4. 填充表格 ==========
    print("\n填充表格...")

    # 表格0: 监测点位安装信息
    if len(tables) > 0 and not site_info.empty and data_point_names:
        print("  填充表格0: 监测点位安装信息...")
        _fill_site_info_table(tables[0], site_info, data_point_names)
        stats["tables_filled"] += 1

    # 表格1: 数据收集率统计
    if len(tables) > 1 and data_point_names:
        print("  填充表格1: 数据收集率统计...")
        _fill_collection_rate_table(tables[1], analysis_data, data_point_names, filter_result_path)
        stats["tables_filled"] += 1

    # 表格2: 日降雨量统计
    if has_rainfall_data and len(tables) > 2 and "日降雨量统计" in analysis_data:
        print("  填充表格2: 日降雨量统计...")
        _fill_rainfall_daily_table(tables[2], analysis_data["日降雨量统计"])
        stats["tables_filled"] += 1

    # 表格3: 场次降雨统计
    if has_rainfall_data and len(tables) > 3 and "场次降雨统计" in analysis_data:
        print("  填充表格3: 场次降雨统计...")
        _fill_rainfall_event_table(tables[3], analysis_data["场次降雨统计"])
        stats["tables_filled"] += 1

    # ========== 5. LLM生成风险分析段落 ==========
    if llm_client and "旱天风险" in analysis_data:
        print("\nLLM生成风险分析段落...")
        llm_writer = LLMReportWriter(llm_client)
        dry_risk_df = analysis_data["旱天风险"]

        # 生成充满度描述
        fullness_idx = _find_paragraph_index(doc, "最大充满度情况如下")
        if fullness_idx >= 0:
            fullness_desc = llm_writer.generate_fullness_description(dry_risk_df, cfg.monitoring_round)
            # 找到后续的编号段落并替换
            for i in range(fullness_idx + 1, min(fullness_idx + 5, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if text.startswith("①") or text.startswith("②") or text.startswith("③") or text.startswith("④"):
                    _replace_paragraph_content(doc.paragraphs[i], "")
            if fullness_desc:
                lines = fullness_desc.split("\n")
                for line in lines[1:]:  # 跳过第一行（标题）
                    if line.strip():
                        _insert_paragraphs_after(doc, fullness_idx, [line.strip()])
                        fullness_idx += 1
                stats["llm_generated"] += 1
                print("  生成充满度描述")

        # 生成溢流风险描述
        overflow_idx = _find_paragraph_index(doc, "溢流风险值情况如下")
        if overflow_idx >= 0:
            overflow_desc = llm_writer.generate_overflow_description(dry_risk_df, cfg.monitoring_round)
            for i in range(overflow_idx + 1, min(overflow_idx + 3, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if text.startswith("①") or text.startswith("②"):
                    _replace_paragraph_content(doc.paragraphs[i], "")
            if overflow_desc:
                lines = overflow_desc.split("\n")
                for line in lines[1:]:
                    if line.strip():
                        _insert_paragraphs_after(doc, overflow_idx, [line.strip()])
                        overflow_idx += 1
                stats["llm_generated"] += 1
                print("  生成溢流风险描述")

        # 生成淤积风险描述
        silting_idx = _find_paragraph_index(doc, "淤积风险情况如下")
        if silting_idx >= 0:
            silting_desc = llm_writer.generate_silting_description(dry_risk_df, cfg.monitoring_round)
            for i in range(silting_idx + 1, min(silting_idx + 4, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if text.startswith("①") or text.startswith("②") or text.startswith("③"):
                    _replace_paragraph_content(doc.paragraphs[i], "")
            if silting_desc:
                lines = silting_desc.split("\n")
                for line in lines[1:]:
                    if line.strip():
                        _insert_paragraphs_after(doc, silting_idx, [line.strip()])
                        silting_idx += 1
                stats["llm_generated"] += 1
                print("  生成淤积风险描述")

        # 生成雨天风险描述
        if has_rainfall_data and "雨天溢流风险" in analysis_data:
            rainy_idx = _find_paragraph_index(doc, "溢流风险情况如下所示")
            if rainy_idx >= 0:
                event_info = {"date": "监测期间", "rainfall": 0, "level": "小雨"}
                if "场次降雨统计" in analysis_data and not analysis_data["场次降雨统计"].empty:
                    first_event = analysis_data["场次降雨统计"].iloc[0]
                    event_info["date"] = str(first_event.get("开始时间", ""))[:10]
                    event_info["rainfall"] = first_event.get("总降雨量(mm)", 0)
                    event_info["level"] = first_event.get("降雨等级", "小雨")

                rainy_desc = llm_writer.generate_rainy_risk_description(
                    analysis_data["雨天溢流风险"], event_info, cfg.monitoring_round
                )
                for i in range(rainy_idx + 1, min(rainy_idx + 3, len(doc.paragraphs))):
                    text = doc.paragraphs[i].text.strip()
                    if text.startswith("①") or text.startswith("②"):
                        _replace_paragraph_content(doc.paragraphs[i], "")
                if rainy_desc:
                    lines = rainy_desc.split("\n")
                    for line in lines[1:]:
                        if line.strip():
                            _insert_paragraphs_after(doc, rainy_idx, [line.strip()])
                            rainy_idx += 1
                    stats["llm_generated"] += 1
                    print("  生成雨天风险描述")

        # 生成风险分析总结
        summary_idx = _find_paragraph_index(doc, "主要结论如下")
        if summary_idx >= 0:
            summary_desc = llm_writer.generate_risk_summary(
                dry_risk_df,
                analysis_data.get("雨天溢流风险"),
                None,
            )
            # 清除旧的总结段落
            for i in range(summary_idx + 1, min(summary_idx + 6, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if text.startswith("（") and "）" in text:
                    _replace_paragraph_content(doc.paragraphs[i], "")
            if summary_desc:
                lines = summary_desc.split("\n")
                for line in lines[1:]:
                    if line.strip():
                        _insert_paragraphs_after(doc, summary_idx, [line.strip()])
                        summary_idx += 1
                stats["llm_generated"] += 1
                print("  生成风险分析总结")

    # ========== 6. 生成特征曲线图 ==========
    img_dir = output_file.parent / "特征曲线图"
    img_dir.mkdir(parents=True, exist_ok=True)

    for point_name, curve in dry_curve_data.items():
        if curve is not None and len(curve) > 0:
            img_path = _generate_pattern_curve_image(curve, point_name, img_dir)
            if img_path:
                stats["images_inserted"] += 1
        stats["points_processed"] += 1

    # ========== 7. 插入特征曲线图 ==========
    if len(tables) > 4 and data_point_names:
        print("\n插入特征曲线图...")
        inserted = _insert_curve_images(doc, img_dir, data_point_names)
        print(f"  插入图片: {inserted} 张")

    # ========== 8. 填充风险表格 ==========
    # 表格17: 旱天风险
    if len(tables) > 17 and "旱天风险" in analysis_data:
        print("填充表格17: 旱天风险...")
        _fill_risk_table(tables[17], analysis_data["旱天风险"])
        stats["tables_filled"] += 1

    # 表格18: 雨天溢流风险
    if has_rainfall_data and len(tables) > 18 and "雨天溢流风险" in analysis_data:
        print("填充表格18: 雨天溢流风险...")
        df = analysis_data["雨天溢流风险"]
        if not df.empty:
            _fill_table_with_df(tables[18], df, start_row=1)
        stats["tables_filled"] += 1

    # ========== 9. 无降雨数据处理 ==========
    if not has_rainfall_data:
        print("删除雨天相关章节...")
        _remove_rainy_sections(doc)

    # ========== 10. 保存报告 ==========
    doc.save(output_file)
    print(f"\n保存报告: {output_file}")

    stats["tables_filled"] = min(stats["tables_filled"], len(tables))

    print(f"\n报告组装完成:")
    print(f"  - 填充表格: {stats['tables_filled']} 个")
    print(f"  - 插入图片: {stats['images_inserted']} 张")
    print(f"  - 处理点位: {stats['points_processed']} 个")
    print(f"  - 文字替换: {stats['text_replaced']} 处")
    print(f"  - LLM生成: {stats['llm_generated']} 段")

    return {
        "output_file": output_file,
        "stats": stats,
    }
