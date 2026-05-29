"""排污规律章节生成模块

从 Excel 读取排污规律分析结果，生成报告中的排污规律章节。
"""

from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .table_manager import adjust_curve_image_tables


def generate_pattern_section(
    doc: Document,
    pattern_df: pd.DataFrame,
    curve_image_dir: Path,
    insert_after_keyword: str = "第一轮监测点排污规律统计",
) -> int:
    """
    生成完整的排污规律分析章节。

    Args:
        doc: Word 文档对象
        pattern_df: 排污规律分析 DataFrame，包含：
            - 点位编号
            - 分类 (1/2/3)
            - 分类名称
            - 排污规律描述
        curve_image_dir: 特征曲线图目录
        insert_after_keyword: 插入位置关键词

    Returns:
        处理的点位数量
    """
    if pattern_df.empty:
        return 0

    # 找到插入位置
    insert_idx = _find_paragraph_index(doc, insert_after_keyword)
    if insert_idx == -1:
        print(f"警告: 未找到插入位置 '{insert_after_keyword}'")
        return 0

    # 按分类分组
    classified = _classify_points(pattern_df)

    # 生成分类统计段落
    summary_text = _build_classification_summary(classified)

    # 删除旧的分类统计段落（从插入位置到第一个分类描述之前）
    _clear_old_content(doc, insert_idx)

    # 插入新的分类统计段落
    _insert_text_after(doc, insert_idx, summary_text)

    # 为每个分类生成描述
    point_names = []
    current_idx = insert_idx + 2  # 跳过标题和统计段落

    for class_id in [1, 2, 3]:
        points = classified.get(class_id, [])
        if not points:
            continue

        # 插入分类标题
        class_title = _get_class_title(class_id)
        _insert_text_after(doc, current_idx, class_title)
        current_idx += 1

        # 为每个点位插入描述和图片
        for point_info in points:
            point_name = point_info["点位编号"]
            description = point_info.get("排污规律描述", "")
            point_names.append(point_name)

            # 插入点位描述
            if description:
                _insert_text_after(doc, current_idx, description)
                current_idx += 1

            # 插入特征曲线图
            img_path = curve_image_dir / f"{point_name}_特征曲线.png"
            if img_path.exists():
                _insert_image_paragraph(doc, current_idx, img_path, point_name)
                current_idx += 1

    return len(point_names)


def _find_paragraph_index(doc: Document, keyword: str) -> int:
    """通过关键词定位段落索引"""
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            return i
    return -1


def _classify_points(pattern_df: pd.DataFrame) -> Dict[int, List[Dict]]:
    """按分类整理点位"""
    result: Dict[int, List[Dict]] = {1: [], 2: [], 3: []}

    for _, row in pattern_df.iterrows():
        class_id = int(row.get("分类", 0))
        if class_id in [1, 2, 3]:
            result[class_id].append(row.to_dict())

    return result


def _build_classification_summary(classified: Dict[int, List[Dict]]) -> str:
    """构建分类统计段落"""
    parts = []

    for class_id in [1, 2, 3]:
        points = classified.get(class_id, [])
        if not points:
            continue

        class_name = _get_class_name(class_id)
        point_names = [p.get("点位编号", "") for p in points]
        count = len(points)

        if class_id == 1:
            parts.append(
                f"流量特征曲线属于第一类符合生活用水规律的有{count}处监测点位，为{'、'.join(point_names)}"
            )
        elif class_id == 2:
            parts.append(
                f"流量特征曲线属于第二类有波峰或波谷但不符合生活用水规律的共有{count}处点位，为{'、'.join(point_names)}"
            )
        else:
            parts.append(
                f"流量特征曲线属于第三类无明显波峰或波谷的共有{count}处点位，为{'、'.join(point_names)}"
            )

    return "；".join(parts) + "。"


def _get_class_name(class_id: int) -> str:
    """获取分类名称"""
    names = {
        1: "符合生活用水规律",
        2: "有波峰或波谷但不符合生活用水规律",
        3: "无明显波峰或波谷",
    }
    return names.get(class_id, "")


def _get_class_title(class_id: int) -> str:
    """获取分类标题"""
    titles = {
        1: "（1）监测点位流量特征曲线符合生活用水规律",
        2: "（2）监测点位流量特征曲线有波峰或波谷但不符合生活用水规律",
        3: "（3）监测点位流量特征曲线无明显波峰或波谷",
    }
    return titles.get(class_id, "")


def _clear_old_content(doc: Document, start_idx: int) -> None:
    """清除旧内容（从指定位置到下一个章节标题）"""
    # 找到下一个章节标题的位置
    end_idx = start_idx + 1
    keywords = ["（1）", "（2）", "（3）", "本章小结", "污水系统运行风险"]

    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        for kw in keywords:
            if text.startswith(kw):
                end_idx = i
                break
        if end_idx != start_idx + 1:
            break

    # 清除段落内容（不删除段落本身）
    for i in range(start_idx + 1, end_idx):
        doc.paragraphs[i].text = ""


def _insert_text_after(doc: Document, after_idx: int, text: str) -> None:
    """在指定段落后插入文本"""
    para = doc.paragraphs[after_idx]

    # 在段落末尾添加新段落
    new_para = para.insert_paragraph_before(text)

    # 移动到正确位置
    para._element.addnext(new_para._element)


def _insert_image_paragraph(
    doc: Document,
    insert_idx: int,
    img_path: Path,
    point_name: str,
) -> None:
    """插入图片段落"""
    para = doc.paragraphs[insert_idx]

    # 清空段落
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加图片
    run = para.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))


def insert_curve_images_to_tables(
    doc: Document,
    curve_image_dir: Path,
    point_names: List[str],
    start_table_idx: int = 4,
) -> int:
    """
    将特征曲线图插入到对应的表格中。

    Args:
        doc: Word 文档对象
        curve_image_dir: 特征曲线图目录
        point_names: 点位名称列表
        start_table_idx: 起始表格索引

    Returns:
        成功插入的图片数量
    """
    inserted = 0
    tables = doc.tables

    for i, point_name in enumerate(point_names):
        table_idx = start_table_idx + i
        if table_idx >= len(tables):
            break

        table = tables[table_idx]
        if len(table.rows) == 0 or len(table.rows[0].cells) < 2:
            continue

        img_path = curve_image_dir / f"{point_name}_特征曲线.png"
        if not img_path.exists():
            continue

        try:
            row = table.rows[0]

            # 左侧单元格：流量曲线图
            cell_a = row.cells[0]
            cell_a.text = ""
            para_a = cell_a.paragraphs[0]
            para_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = para_a.add_run()
            run_a.add_picture(str(img_path), width=Inches(2.8))

            # 右侧单元格：液位曲线图（同一张图）
            cell_b = row.cells[1]
            cell_b.text = ""
            para_b = cell_b.paragraphs[0]
            para_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_b = para_b.add_run()
            run_b.add_picture(str(img_path), width=Inches(2.8))

            inserted += 2

        except Exception as e:
            print(f"插入图片失败 {point_name}: {e}")

    return inserted


def get_pattern_point_names(pattern_df: pd.DataFrame) -> List[str]:
    """从排污规律分析结果提取点位名称列表"""
    if pattern_df.empty:
        return []

    if "点位编号" in pattern_df.columns:
        return pattern_df["点位编号"].dropna().tolist()

    return []
